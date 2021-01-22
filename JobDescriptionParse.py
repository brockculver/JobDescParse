import nltk
from nltk import corpus
from nltk.corpus import wordnet as wn
import docx
import time
import os

class JDparser:
    def readDesc(self,name,type):
        if(type=="docx"):
            doc = docx.Document(name+"."+type)
            i=0
            jobDescStr=""
            for p in doc.paragraphs:
                i+=1
            for x in range(0, i):
                jobDescStr=jobDescStr+doc.paragraphs[x].text.lower()+" "
            return jobDescStr
        else:
            #TODO: insert support for other file formats (.pdf, .txt, etc.)
            return


    def degReq(self,desc):
        edInfo=edR()
        degReq=""

        idx=edInfo.matchDegree(desc)
        if idx[0]!=-1:
            majReq=desc[idx[0]-20:idx[0]+40]
            # print(idx[1])
            return [idx[1],majReq]
        else:
            majReq="No Matching degree :("
        return [majReq]


    def majReq(self,desc):
        edInfo=edR()
        degReq=""

        idx=edInfo.matchMajor(desc)
        if idx[0]!=-1:
            return idx
            # majReq=desc[idx[0]-20:idx[0]+40]
            # # print(idx[1])
            # return [idx[1],majReq]
        else:
            majReq="No Matching major :("
        return [majReq]


    def getCommonWords(self, desc):
        mostCommon=["Most Common: "]
        allWords = [word for word in desc if len(word) >3]

        allWordExceptStopDist = nltk.FreqDist(allWords)
        mostCommon= mostCommon+allWordExceptStopDist.most_common(10)

        return mostCommon


    def getShortestPathWords(self, desc):
        failedWords=[]
        shortestPathWords=[]

        for w in desc:
            try:
                if wn.synsets(w)[0].shortest_path_distance(wn.synset("skill.n.01"))> 2 and wn.synsets(w)[0].shortest_path_distance(wn.synset("skill.n.01"))<8:
                    shortestPathWords.append(w)
            except:
                    failedWords.append(w)
        shortestPathWords=set(shortestPathWords)
        shortestPathWords=["ShortestPathWords: "]+list(shortestPathWords)
        return shortestPathWords


    def getMajorWords(self, desc, major):
        edInfo=edR()
        stopLoop=0
        majorWords=[]
        for i in major:
            if (i =='computer science' or i =='data science' or i == 'information systems' or i =='engineering' or i =='math') and stopLoop==0:
                stopLoop=1
                for w in range(len(desc)):
                    for z in range(len(edInfo.skillsByMajorCS)):
                        skillSet = edInfo.skillsByMajorCS[z].split(" ")
                        for x in range(len(skillSet)):
                            # print(skillSet[x]," needs to match ",allWords[w+x])
                            if skillSet[x]!=desc[w+x]:
                                # print("didnt match")
                                break
                            if x == len(skillSet)-1:
                                majorWords.append(edInfo.skillsByMajorCS[z])
                                # print("match skill by cs keyword:", edInfo.skillsByMajorCS[z])
        majorWords=set(majorWords)
        majorWords=["Major Words: "]+list(majorWords)
        return majorWords


    def getSoftSkills(self,desc):
        edInfo=edR()
        softSkills=[]
        for w in range(len(desc)):
            for z in range(len(edInfo.softSkillSet)):
                skillSet = edInfo.softSkillSet[z].split(" ")
                for x in range(len(skillSet)):
                    if skillSet[x]!=desc[w+x]:
                        break
                    if x == len(skillSet)-1:
                        softSkills.append(edInfo.softSkillSet[z])

        softSkills=set(softSkills)
        softSkills=["Soft Skills: "]+list(softSkills)
        return softSkills

    def getSkillWords(self, desc, major):
        allWords = nltk.tokenize.word_tokenize(desc)
        stopwords = nltk.corpus.stopwords.words('english')
        words = nltk.corpus.words.words()

        allWords = [word for word in allWords if word not in stopwords]

        commonWords=self.getCommonWords(allWords)
        shortPathWords=self.getShortestPathWords(allWords)
        majorWords=self.getMajorWords(allWords,major)
        softWords=self.getSoftSkills(allWords)

        return [commonWords,shortPathWords,majorWords,softWords]

    def mkAllPhrases(self,skillWords):
        phrases=[]

        phrases.append(self.mkMajorPhrases(skillWords[2]))
        phrases.append(self.mkSoftPhrases(skillWords[3]))
        phrases.append(self.mkMostCommonPhrases(skillWords[0]))
        phrases.append(self.mkSkillPhrases(skillWords[1]))
        return phrases

    def mkSkillPhrases(self,skills):
        skills.remove(skills[0])
        skillz=''
        for x in range(len(skills)):
            skillz=skillz+skills[x]+', '
            if x %10==0:
                skillz=skillz+"\n"
        skillz=skillz+"\n"
        return skillz

    def mkMostCommonPhrases(self,skills):
        skills.remove(skills[0])
        mstcmn=''
        for x in range(len(skills)):
            mstcmn=mstcmn+skills[x][0]+', '
        mstcmn=mstcmn+'\n'
        return mstcmn

    def mkMajorPhrases(self, skills):
        try:
            count=0

            skills.remove(skills[0])
            skills.sort(key=len)
            skillsToDelete=[]

            for s in range(len(skills)):
                if s<(len(skills)-1):
                    for i in range(s+1, len(skills)):
                        split2nd=skills[i].split(' ')
                        for j in range(len(split2nd)):
                            if skills[s]==split2nd[j]:
                                skillsToDelete.append(skills[s])
                                break

            for s in skillsToDelete:
                skills.remove(s)
            if (len(skills))>1:
                majorPhrase="Additional technical skills include "
                if (len(skills))>7:
                    majorPhrase=majorPhrase+": "
                for word in skills:
                    count=count+1
                    majorPhrase=majorPhrase+word+", "
                    if count %10==0:
                        majorPhrase=majorPhrase+"\n"
                majorPhrase=majorPhrase+'and other competencies.'
            else:
                majorPhrase="Additionally, technical proficiency with "+skills[0]

            majorPhrase=majorPhrase+'\n'
        except:
            majorPhrase='No major match or no technical skills\n'

        return majorPhrase

    def mkSoftPhrases(self, skills):
        count=0
        skills.remove(skills[0])
        skills.sort(key=len)
        skillsToDelete=[]
        for s in range(len(skills)):
            if s<(len(skills)-1):
                for i in range(s+1, len(skills)):
                    split2nd=skills[i].split(' ')
                    for j in range(len(split2nd)):
                        if skills[s]==split2nd[j]:
                            skillsToDelete.append(skills[s])
                            break
        for s in skillsToDelete:
            skills.remove(s)
        if (len(skills))>1:
            softPhrase="Strong soft skills include "
            if (len(skills))>7:
                softPhrase=softPhrase+": "
            for word in skills:
                count=count+1
                softPhrase=softPhrase+word+", "
                if count %10==0:
                    softPhrase=softPhrase+"\n"
            softPhrase=softPhrase+'and others.'
        else:
            softPhrase="Often described as "+skills[0]

        softPhrase=softPhrase+'\n'
        return softPhrase


    def getMajorPhrase(self):
        major=self.majReq(self.jobDesc)
        skills=self.getSkillWords(self.jobDesc, major)

    def get_attributes(self):
        education=self.degReq(self.jobDesc)
        major=self.majReq(self.jobDesc)
        skills=self.getSkillWords(self.jobDesc, major)
        return skills


    def __init__(self, DocName):
       self.DocName = DocName.split(".")[0]
       self.DocType=DocName.split(".")[1]
       self.jobDesc=self.readDesc(self.DocName,self.DocType)



class edR:
    edDegrees=[ "master", "associate", "bachelor", "b.s", "m.s", "phd","dr.","doctor", "degree","high school"]

    edMjrs=["""geological and geophysical engineering""","""industrial and manufacturing engineering""","""materials engineering""","""materials science""","""mechanical engineering""","""metallurgical engineering""","""mining and mineral engineering""","""naval architecture and marine engineering""","""nuclear engineering""","""petroleum engineering""","""miscellaneous engineering""","""engineering technologies""","""engineering and industrial management""","""electrical engineering technology""","""industrial production technologies""","""mechanical engineering related technologies""","""miscellaneous engineering technologies""","""materials science""","""nutrition sciences""","""general medical and health services""","""communication disorders sciences and services""","""health and medical administrative services""","""medical assisting services""","""medical technologies technicians""","""health and medical preparatory programs""","""nursing""","""pharmacy pharmaceutical sciences and administration""","""treatment therapy professions""","""community and public health""","""miscellaneous health medical professions""","""area ethnic and civilization studies""","""linguistics and comparative language and literature""","""french german latin and other common foreign language studies""","""other foreign languages""","""english language and literature""","""composition and rhetoric""","""liberal arts""","""humanities""","""intercultural and international studies""","""philosophy""","""religious studies""","""theology""","""religious vocations""","""anthropology""","""archeology""","""art history and criticism""","""history""","""united states history""","""cosmetology services and culinary arts""","""family and consumer sciences""","""military technologies""","""physical fitness parks recreation and leisure""","""construction services""","""electrical, mechanical, and precision technologies and production""","""transportation sciences and technologies""","""multi/interdisciplinary studies""","""court reporting""","""pre-law""","""legal studies""","""criminal justice""","""fire protection""","""public administration""","""public policy""","""n/a (less than bachelor's degree)""","""physical sciences""","""astronomy and astrophysics""","""atmospheric sciences and meteorology""","""chemistry""","""geology and earth science""","""geosciences""","""oceanography""","""physics""","""multi-disciplinary or general science""","""nuclear, industrial radiology, and biological technologies""","""psychology""","""educational psychology""","""clinical psychology""","""counseling psychology""","""industrial and organizational psychology""","""social psychology""","""miscellaneous psychology""","""human services and community organization""","""social work""","""interdisciplinary social sciences""","""general social sciences""","""economics""","""criminology""","""geography""","""international relations""","""political science and government""","""sociology""","""miscellaneous social sciences""","""environmental engineering""","""general agriculture""","""agriculture production and management""","""agricultural economics""","""animal sciences""","""food science""","""plant science and agronomy""","""soil science""","""miscellaneous agriculture""","""forestry""","""natural resources management""","""fine arts""","""drama and theater arts""","""music""","""visual and performing arts""","""commercial art and graphic design""","""film video and photographic arts""","""studio arts""","""fine arts""","""environmental science""","""biology""","""biochemical sciences""","""botany""","""molecular biology""","""ecology""","""genetics""","""microbiology""","""pharmacology""","""physiology""","""zoology""","""neuroscience""","""miscellaneous biology""","""cognitive science and biopsychology""","""general business""","""accounting""","""actuarial science""","""business management and administration""","""operations logistics and e-commerce""","""business economics""","""marketing and marketing research""","""finance""","""human resources and personnel management""","""international business""","""hospitality management""","""management information systems and statistics""","""miscellaneous business & medical administration""","""degree in communications""","""communications degree""","""journalism""","""mass media""","""advertising and public relations""","""communication technologies""","""information systems""","""computer programming and data processing""","""computer science""","""information sciences""","""computer administration management and security""","""computer networking and telecommunications""","""mathematics""","""applied mathematics""","""statistics and decision science""","""math""","""general education""","""educational administration and supervision""","""school student counseling""","""elementary education""","""mathematics teacher education""","""physical and health education teaching""","""early childhood education""","""science and computer teacher education""","""secondary teacher education""","""special needs education""","""social science or history teacher education""","""teacher education: multiple levels""","""language and drama education""","""art and music education""","""miscellaneous education""","""library science""","""architecture""","""engineering""","""aerospace engineering""","""biological engineering""","""architectural engineering""","""biomedical engineering""","""chemical engineering""","""civil engineering""","""computer engineering""","""electrical engineering""","""engineering mechanics physics and science""", """data science"""]

    softSkillSet=['trustworthiness','honest','kind','ethical','moral','focusing','attentiveness','displaying personal values','having principles','integrity','courtesy','business etiquette','manners','polite','graciousness','patience','apologizing','respectfulness','phone etiquette','gratefulness','decency','humour','nurturing','empathy','self control','sociability','showing warmth','personability','consideration','friendliness','body language','social','communication','public speaking','clarity of speech and writing','non-verbal communication','presentation','listening','open-mindedness','giving and receiving feedback','empathy','choosing a communication medium','knowing when to communicate','open mindedness','flexibility','flexible','willingness to change','adaptability','lifelong learning','teachability','acceptance','adjustability','versatility','improvisation','calmness','solutions','teamwork','cooperation','agreeability','helpfulness','supportiveness','collaboration','likeability','persuasiveness','influence','conflict resolution','facilitation','responsibility','follow through','conscientious','accountable','aspiring','virtuous','resourcefulness','common sense','dependability','maturity','ethical','innovative','strong work ethic','attitude','happiness','confidence','optimism','enthusiasm','encouraging','courageousness','spiritedness','pluckiness','intrepidity','fortitude','professional','sophistication','credibility','valor','respectability','consideration','pragmatism','leadership','mentoring','conflict management','gallantry','time management','loyal','adaptability','respect','superiority','excellence','driven','analytical','interpersonal','intrapersonal','context','strategic thinking','analytical','context','futuristic','ideation','input','intellection','learner','strategic','relationship building','adaptability','entrpreneurial','quantitative','connectedness','developer','empathy','harmony','includer','individualization','positivity','relator','influencing','activator','command','communication','competition','maximizer','self-assurance','significance','woo','executing','achiever','arranger','belief','consistency','deliberative','discipline','focus','responsibility','security clearance','restorative','passion','microsoft word', 'microsoft excel', 'microsoft powerpoint','excel','powerpoint','microsoft office']

    preTechnicalSkillPhraseList=["Familiarity with ", "Coursework involving ", "Projects worked with ", "Developed knowledge relating to ", "Skills including "]

    skillsByMajorCS=["version control","git","github","java","python","c","open source","mongodb","c++","css","html","ruby","javascript","node.js","node","frontend","backend","full stack","mobile application","application","ajax",".net","react","vue","json","angular","xml","docker","django","gcp","azure","numpy","pandas","tensorflow","aws","pytorch","kubernetes","haskell","rust","golang","r","vba","visual basic","basic","c#","php","groovy","assembly","sql","swift","sas","matlab","julia","dart","cobol","scala","lisp","scripting","artificial intelligence","machine learning","information security","big data","data analysis","analysis","data science","data","hadoop","devops","data mining","spark","oop","object oriented","agile","statistics","structures","systems", "algorithms","concurrent","design","encryption","cloud","deep learning","parallel","networking","database","compiler","calculus","math",'optimization','simulation','debugging','ide','eclipse','intellij','multi-threaded',"linux",'unix',"windows","macos","android",'api',"web development","web","cybersecurity","blockchain",'debug',"development"]

    def matchDegree(self, desc):
        locDegree=-1
        desc=desc.lower()
        for x in edR.edDegrees:
            try:
                # print(x)
                locDegree=desc.find(x)
                #print("locDegreeee: ")
                #print(locDegree)
            finally:
                #print("\n")
                if locDegree!=-1:
                    #print("\n found loc Degree")
                    return [locDegree,x]
        return [locDegree]

    def matchMajor(self, desc):
        anyMatches=0
        hitMajors=[]
        locDegree=-1
        desc=desc.lower()

        for x in edR.edMjrs:
            locDegree=desc.find(x)
            if locDegree!=-1:
                anyMatches+=1
                hitMajors.append(x)
                hitMajors.append(locDegree)
        if anyMatches!=0:
            return hitMajors
        return [locDegree]
